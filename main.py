import os
import json
import re
from google.genai import types
from dotenv import load_dotenv
from google import genai

load_dotenv()

api_key = os.getenv("GEMINI_API_KEY")

if not api_key:
    print("Error: Gemini API key not found.")
    exit()

print("Gemini API key loaded successfully.")

client = genai.Client(api_key=api_key)

from pathlib import Path
from docx import Document


def read_txt(file_path):
    """Read text from a .txt file."""
    with open(file_path, "r", encoding="utf-8") as file:
        return file.read()


def read_docx(file_path):
    """Read text from a .docx file, including tables."""

    try:
        document = Document(file_path)
    except PermissionError:
        print(f"Error: Permission denied while reading '{file_path}'.")
        print("Close the file if it is open in Microsoft Word and try again.")
        return None
    except Exception as error:
        print(f"Error reading DOCX file: {error}")
        return None

    paragraphs = []

    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            paragraphs.append(paragraph.text)

    # Also extract text from tables to make text extraction robust
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                # Append each unique cell block
                if cell_text and cell_text not in paragraphs:
                    paragraphs.append(cell_text)

    return "\n".join(paragraphs)


def read_resume(file_path):
    """Read a resume from either .txt or .docx."""

    path = Path(file_path)

    if not path.exists():
        print(f"Error: {path.name} not found.")
        return None

    if path.suffix.lower() == ".txt":
        return read_txt(path)

    elif path.suffix.lower() == ".docx":
        return read_docx(path)

    else:
        print("Error: Only .txt and .docx files are supported.")
        return None


def validate_resume(resume):
    """Check whether the resume has enough usable content."""

    if not resume.strip():
        print("Error: Resume is empty.")
        return False

    if len(resume.strip()) < 50:
        print("Error: Resume is too short.")
        return False

    return True


def clean_resume(resume):
    """Remove unnecessary spaces and blank lines."""

    lines = resume.splitlines()

    cleaned_lines = []

    for line in lines:
        line = line.strip()

        if line:
            cleaned_lines.append(line)

    return "\n".join(cleaned_lines)

def remove_empty_section(html, section_id):
    """Remove a portfolio section and its navigation link when empty."""

    pattern = rf'\s*<section id="{section_id}".*?</section>'
    html = re.sub(pattern, '', html, flags=re.DOTALL)

    nav_pattern = rf'\s*<a href="#{section_id}">.*?</a>'
    html = re.sub(nav_pattern, '', html)

    return html

def render_portfolio(portfolio_data):
    """Convert Gemini portfolio data into HTML and save it."""

    with open("template.html", "r", encoding="utf-8") as file:
        template = file.read()

    # Load style.css content for inline styling
    try:
        with open("style.css", "r", encoding="utf-8") as css_file:
            style_css = css_file.read()
    except Exception as e:
        print(f"Warning: Could not read style.css for inlining: {e}")
        style_css = "/* Failed to load style.css */"

    # Inline style.css
    html = template.replace("{{ style_css }}", style_css)

    # Basic information
    html = html.replace(
        "{{ full_name }}",
        portfolio_data["full_name"]
    )

    html = html.replace(
        "{{ headline }}",
        portfolio_data["headline"]
    )
    if portfolio_data["projects"]:
        project_button = """
    <a href="#projects" class="hero-button">
        View My Projects
    </a>
    """
    else:
        project_button = ""

    html = html.replace("{{ project_button }}", project_button)

    html = html.replace(
        "{{ about }}",
        portfolio_data["about"]
    )

    # Education
    education_html = ""

    for education in portfolio_data["education"]:
        education_html += f"""
        <div class="education-card">
           <h3>{education["degree"]}</h3>
        """

        if education["institution"].strip():
                 education_html += f"""
            <p>{education["institution"]}</p>
            """

        if education["year"].strip():
            education_html += f"""
            <p>{education["year"]}</p>
            """

        if education["details"].strip():
            education_html += f"""
            <p>{education["details"]}</p>
            """

        education_html += """
        </div>
        """

    html = html.replace("{{ education }}", education_html)

    # Skills
    skills_html = ""

    for skill in portfolio_data["skills"]:
        skills_html += f"""
        <span class="skill">{skill}</span>
        """

    html = html.replace("{{ skills }}", skills_html)

    # Projects
    projects_html = ""

    for project in portfolio_data["projects"]:
        projects_html += f"""
        <div class="project-card">
            <h3>{project["name"]}</h3>
        """

        if project["description"].strip():
            projects_html += f"""
           <p>{project["description"]}</p>
            """

        if project["technologies"]:
            technologies = ", ".join(project["technologies"])
            projects_html += f"""
            <p><strong>Technologies:</strong> {technologies}</p>
            """

        if project["link"]:
            projects_html += f"""
            <a href="{project["link"]}" target="_blank">
                View Project
            </a>
            """

        projects_html += """
        </div>
        """

    html = html.replace("{{ projects }}", projects_html)

    # Experience
    experience_html = ""

    for exp in portfolio_data["experience"]:
        experience_html += f"""
    <div class="experience-card">
        <h3>{exp["role"]}</h3>
    """

        if exp["company"].strip():
            experience_html += f"""
            <p>{exp["company"]}</p>
            """

        if exp["duration"].strip():
            experience_html += f"""
            <p>{exp["duration"]}</p>
            """

        if exp["description"].strip():
            experience_html += f"""
            <p>{exp["description"]}</p>
            """

        experience_html += """
        </div>
        """

    html = html.replace("{{ experience }}", experience_html)

    # Achievements
    achievements_html = ""

    for achievement in portfolio_data["achievements"]:
        achievements_html += f"""
        <div class="achievement-card">
            <p>{achievement}</p>
        </div>
        """

    html = html.replace("{{ achievements }}", achievements_html)

    # Contact
    contact = portfolio_data["contact"]

    contact_html = ""

    if contact["email"]:
        contact_html += f"""
        <p>
            Email:
            <a href="mailto:{contact["email"]}">
                {contact["email"]}
            </a>
        </p>
        """

    if contact["phone"]:
        contact_html += f"""
        <p>Phone: {contact["phone"]}</p>
        """

    if contact["linkedin"]:
        contact_html += f"""
        <p>
            <a href="{contact["linkedin"]}" target="_blank">
                LinkedIn
            </a>
        </p>
        """

    if contact["github"]:
        contact_html += f"""
        <p>
            <a href="{contact["github"]}" target="_blank">
                GitHub
            </a>
        </p>
        """

    if contact["website"]:
        contact_html += f"""
        <p>
            <a href="{contact["website"]}" target="_blank">
                Website
            </a>
        </p>
        """

    html = html.replace("{{ contact }}", contact_html)

    # Remove sections that have no information
    if not portfolio_data["about"].strip():
        html = remove_empty_section(html, "about")
    if not portfolio_data["education"]:
        html = remove_empty_section(html, "education")

    if not portfolio_data["skills"]:
        html = remove_empty_section(html, "skills")

    if not portfolio_data["projects"]:
        html = remove_empty_section(html, "projects")

    if not portfolio_data["experience"]:
        html = remove_empty_section(html, "experience")

    if not portfolio_data["achievements"]:
        html = remove_empty_section(html, "achievements")

    contact = portfolio_data["contact"]

    if not any(contact.values()):
        html = remove_empty_section(html, "contact")

    # Save generated portfolio
    with open("portfolio.html", "w", encoding="utf-8") as file:
        file.write(html)

    print("\nPortfolio generated successfully!")
    print("Saved as: portfolio.html")

def choose_resume_file():
    """Choose a resume from supported TXT or DOCX files."""

    txt_file = Path("resume.txt")
    docx_file = Path("resume.docx")

    txt_exists = txt_file.exists()
    docx_exists = docx_file.exists()

    if txt_exists and docx_exists:
        print("\nBoth resume.txt and resume.docx were found.")
        choice = input("Enter the resume filename to use: ").strip()

        if choice not in ("resume.txt", "resume.docx"):
            print("Error: Please enter resume.txt or resume.docx.")
            return None

        return choice

    if txt_exists:
        print("\nUsing resume.txt")
        return "resume.txt"

    if docx_exists:
        print("\nUsing resume.docx")
        return "resume.docx"

    print("\nError: No resume found.")
    print("Please add either resume.txt or resume.docx.")
    return None


def main():
    """Read, validate, and clean the resume."""

    resume_file = choose_resume_file()

    if resume_file is None:
        return

    resume_text = read_resume(resume_file)

    if resume_text is None:
        return

    if not validate_resume(resume_text):
        return

    cleaned_resume = clean_resume(resume_text)

    print("\n----- CLEANED RESUME -----\n")
    print(cleaned_resume)

    portfolio_data = generate_portfolio_data(cleaned_resume)

    if portfolio_data is None:
        return

    print("\n----- GEMINI PORTFOLIO DATA -----")
    print(json.dumps(portfolio_data, indent=4))

    render_portfolio(portfolio_data)

def generate_portfolio_data(resume_text):
    """Send the cleaned resume to Gemini and get structured portfolio data."""

    portfolio_schema = {
        "type": "object",
        "properties": {
            "full_name": {
                "type": "string"
            },
            "headline": {
                "type": "string"
            },
            "about": {
                "type": "string"
            },
            "education": {
                "type": "array",
                "items": {
                    "type": "object",
                    "properties": {
                        "degree": {"type": "string"},
                        "institution": {"type": "string"},
                        "year": {"type": "string"},
                        "details": {"type": "string"}
                    },
                    "required": [
                        "degree",
                        "institution",
                        "year",
                        "details"
                    ]
                }
            },
            "skills": {
                "type": "array",
                "items": {
                    "type": "string"
                }
            },
            "projects": {
                "type": "array",
                "items": {
                    "type": "object",
                    "properties": {
                        "name": {"type": "string"},
                        "description": {"type": "string"},
                        "technologies": {
                            "type": "array",
                            "items": {"type": "string"}
                        },
                        "link": {"type": "string"}
                    },
                    "required": [
                        "name",
                        "description",
                        "technologies",
                        "link"
                    ]
                }
            },
            "experience": {
                "type": "array",
                "items": {
                    "type": "object",
                    "properties": {
                        "role": {"type": "string"},
                        "company": {"type": "string"},
                        "duration": {"type": "string"},
                        "description": {"type": "string"}
                    },
                    "required": [
                        "role",
                        "company",
                        "duration",
                        "description"
                    ]
                }
            },
            "achievements": {
    "type": "array",
    "items": {
        "type": "string"
    }
},
            "contact": {
                "type": "object",
                "properties": {
                    "email": {"type": "string"},
                    "phone": {"type": "string"},
                    "linkedin": {"type": "string"},
                    "github": {"type": "string"},
                    "website": {"type": "string"}
                },
                "required": [
                    "email",
                    "phone",
                    "linkedin",
                    "github",
                    "website"
                ]
            }
        },
        "required": [
    "full_name",
    "headline",
    "about",
    "education",
    "skills",
    "projects",
    "experience",
    "achievements",
    "contact"
]
    }

    prompt = f"""
You are a resume information extraction system.

Read the resume below and extract information for a personal portfolio.

IMPORTANT RULES:
1. Use ONLY information explicitly present in the resume.
2. Do NOT invent skills, projects, companies, dates, links, education,
   experience, achievements, or contact information.
3. If a text field is not available, return an empty string.
4. If a list section is not available, return an empty list.
5. Keep the information truthful to the resume.
6. The "about" section should be a concise professional summary based
   only on information present in the resume.
7. Extract awards, certifications, achievements, or notable results
   into the "achievements" section.
8. Return JSON matching the provided schema.

RESUME:
----------------
{resume_text}
----------------
"""

    models_to_try = ["gemini-3.6-flash", "gemini-3.5-flash", "gemini-3.7-flash"]
    last_error = None
    response = None

    for model_name in models_to_try:
        try:
            print(f"Attempting to generate portfolio data using model: {model_name}...")
            response = client.models.generate_content(
                model=model_name,
                contents=prompt,
                config=types.GenerateContentConfig(
                    response_mime_type="application/json",
                    response_schema=portfolio_schema
                )
            )
            # If successful, break
            print(f"Successfully generated portfolio data using model: {model_name}.")
            break
        except Exception as e:
            last_error = e
            print(f"Model {model_name} failed: {e}")
            continue

    if response is None:
        print("\nError: Gemini API request failed.")
        print("Please check your API key or try again later.")
        return None

    try:
        return json.loads(response.text)
    except json.JSONDecodeError:
        print("\nError: Gemini returned an invalid JSON response.")
        print("Please try again.")
        return None

if __name__ == "__main__":
    main()